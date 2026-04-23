from __future__ import annotations

import html
import io
import logging
import os
import smtplib
import ssl
import sys
import time
import uuid
from collections import defaultdict, deque
from email.message import EmailMessage
from email.utils import formataddr
from typing import Literal

import certifi
import fitz
import pytesseract
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from openpyxl import Workbook
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter
from PIL import Image
from pydantic import BaseModel, EmailStr, Field, field_validator
from pypdf import PdfReader, PdfWriter

_BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
if _BACKEND_DIR not in sys.path:
    sys.path.insert(0, _BACKEND_DIR)

from services.cloudflare_ai import (
    GENERATED_IMAGES_DIR,
    build_image_provider,
    ensure_generated_images_dir,
)
from utils.prompt_enhancer import enhance_prompt

_FRONTEND_ROOT = os.path.dirname(_BACKEND_DIR)
_REPO_ROOT = os.path.dirname(_FRONTEND_ROOT)
for _dotenv_path in (
    os.path.join(_REPO_ROOT, ".env"),
    os.path.join(_FRONTEND_ROOT, ".env"),
    os.path.join(_BACKEND_DIR, ".env"),
):
    if os.path.isfile(_dotenv_path):
        load_dotenv(_dotenv_path, override=False)

app = FastAPI(title="Converter Backend", version="1.0.0")
logging.basicConfig(
    level=os.environ.get("LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
)
logger = logging.getLogger(__name__)
generated_images_dir = ensure_generated_images_dir(_BACKEND_DIR)

_cors_raw = os.environ.get("CORS_ALLOW_ORIGINS", "*").strip()
_cors_origins = [o.strip() for o in _cors_raw.split(",") if o.strip()] or ["*"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.mount(
    f"/{GENERATED_IMAGES_DIR}",
    StaticFiles(directory=str(generated_images_dir)),
    name="generated_images",
)


def _stream_bytes(data: bytes, filename: str, media_type: str):
    headers = {"x-filename": filename}
    return StreamingResponse(io.BytesIO(data), media_type=media_type, headers=headers)


def _safe_name(name: str, fallback: str) -> str:
    clean = (name or fallback).strip().replace("/", "_").replace("\\", "_")
    return clean or fallback


def _read_image(upload: UploadFile) -> Image.Image:
    raw = upload.file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Empty file.")
    try:
        img = Image.open(io.BytesIO(raw))
        img.load()
        return img
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=400, detail="Invalid image file.") from exc


def _is_pdf_bytes(data: bytes) -> bool:
    return len(data) >= 4 and data[:4] == b"%PDF"


def _resize_for_ocr(img: Image.Image, max_side: int = 2000) -> Image.Image:
    w, h = img.size
    if max(w, h) <= max_side:
        return img
    scale = max_side / max(w, h)
    new_w = max(1, int(w * scale))
    new_h = max(1, int(h * scale))
    return img.resize((new_w, new_h), Image.Resampling.LANCZOS)


def _image_to_rgb(img: Image.Image) -> Image.Image:
    if img.mode in ("RGBA", "P", "LA"):
        return img.convert("RGB")
    if img.mode != "RGB":
        return img.convert("RGB")
    return img


def _extract_text_from_docx(raw: bytes) -> str:
    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = " | ".join((c.text or "").strip() for c in row.cells)
            if cells.strip():
                parts.append(cells)
    return "\n".join(parts).strip()


# Minimum chars to treat PDF as "digital" (embedded text) and skip raster OCR
_MIN_EMBEDDED_PDF_TEXT = 60

# Lightweight OCR: small Python package (pytesseract) + system Tesseract binary only.
# No PyTorch / heavy ML stacks.


def _ocr_image_pil(img: Image.Image, language: str) -> str:
    rgb = _image_to_rgb(img)
    try:
        return (pytesseract.image_to_string(rgb, lang=language) or "").strip()
    except pytesseract.TesseractNotFoundError as exc:
        raise HTTPException(
            status_code=503,
            detail=(
                "Tesseract OCR binary is not installed. Lightweight setup: "
                "macOS: brew install tesseract  (optional langs: brew install tesseract-lang). "
                "Then restart the API. Python side stays light (only pytesseract)."
            ),
        ) from exc


class ContactPayload(BaseModel):
    name: str = Field(min_length=2, max_length=120)
    email: EmailStr
    subject: str = Field(default="", max_length=200)
    message: str = Field(min_length=6, max_length=8000)

    @field_validator("name", "subject", "message", mode="before")
    @classmethod
    def _strip_text(cls, v: object) -> object:
        return v.strip() if isinstance(v, str) else v

    @field_validator("email", mode="before")
    @classmethod
    def _strip_email(cls, v: object) -> object:
        return v.strip() if isinstance(v, str) else v


class ImageGenerationRequest(BaseModel):
    prompt: str = Field(min_length=5, max_length=700)

    @field_validator("prompt", mode="before")
    @classmethod
    def _strip_prompt(cls, v: object) -> object:
        return v.strip() if isinstance(v, str) else v


class ImageGenerationResponse(BaseModel):
    message: str
    image_url: str


class IpRateLimiter:
    def __init__(self, *, max_requests: int, window_seconds: int) -> None:
        self.max_requests = max_requests
        self.window_seconds = window_seconds
        self._hits: dict[str, deque[float]] = defaultdict(deque)

    def check(self, client_ip: str) -> None:
        now = time.time()
        bucket = self._hits[client_ip]
        cutoff = now - self.window_seconds
        while bucket and bucket[0] < cutoff:
            bucket.popleft()

        if len(bucket) >= self.max_requests:
            raise HTTPException(
                status_code=429,
                detail="Rate limit exceeded. Max 5 requests per minute.",
            )
        bucket.append(now)


rate_limiter = IpRateLimiter(max_requests=5, window_seconds=60)


def _reload_ai_env() -> None:
    # Never override platform env vars (e.g. Railway secrets).
    # Load from .env only as a fallback for local development.
    for path in (
        os.path.join(_REPO_ROOT, ".env"),
        os.path.join(_FRONTEND_ROOT, ".env"),
        os.path.join(_BACKEND_DIR, ".env"),
    ):
        if os.path.isfile(path):
            load_dotenv(path, override=False)


def _env(key: str, default: str = "") -> str:
    return (os.environ.get(key) or default).strip().strip("'\"")


def _contact_from_header() -> str:
    raw = _env("DEFAULT_FROM_EMAIL") or _env("MAIL_FROM")
    inbox = _env("CONTACT_INBOX_EMAIL")
    from_name = _env("DEFAULT_FROM_NAME") or "Sohail"

    if raw and "@" in raw and "<" not in raw:
        return raw
    if raw and "<" in raw and ">" in raw:
        inner = raw[raw.index("<") + 1 : raw.index(">")].strip()
        if "@" in inner:
            return raw
        if inbox and "@" in inbox:
            return formataddr((from_name, inbox))
    if inbox and "@" in inbox:
        return formataddr((from_name, inbox))
    raise RuntimeError(
        "Set CONTACT_INBOX_EMAIL and DEFAULT_FROM_EMAIL (SendGrid-verified), e.g. Sohail <you@domain.com>."
    )


def _contact_inbox() -> str:
    explicit = _env("CONTACT_INBOX_EMAIL")
    if explicit and "@" in explicit:
        return explicit
    from_h = _contact_from_header()
    if "<" in from_h and ">" in from_h:
        inner = from_h[from_h.index("<") + 1 : from_h.index(">")].strip()
        if "@" in inner:
            return inner
    if "@" in from_h:
        return from_h
    raise RuntimeError("Set CONTACT_INBOX_EMAIL.")


def _send_smtp(messages: list[EmailMessage]) -> None:
    host = _env("EMAIL_HOST", "smtp.sendgrid.net")
    port = int(_env("EMAIL_PORT", "587") or "587")
    user = _env("SENDGRID_USERNAME", "apikey")
    password = _env("SENDGRID_PASSWORD") or _env("SENDGRID_API_KEY")
    if not password:
        raise RuntimeError("Set SENDGRID_PASSWORD or SENDGRID_API_KEY.")

    if _env("EMAIL_TLS_VERIFY", "true").lower() in ("0", "false", "no", "off"):
        ctx = ssl.SSLContext(ssl.PROTOCOL_TLS_CLIENT)
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE
    else:
        ctx = ssl.create_default_context(cafile=certifi.where())

    with smtplib.SMTP(host, port, timeout=30) as smtp:
        smtp.starttls(context=ctx)
        smtp.login(user, password)
        for msg in messages:
            smtp.send_message(msg)


def _contact_template_plain(name: str, message: str) -> str:
    wave, pin, smile = "\U0001f64c", "\U0001f4cc", "\U0001f60a"
    return (
        f"Hi {name},\n\n"
        f"Thanks for reaching out to us! {wave}\n\n"
        "We've received your message and our team will get back to you as soon as possible.\n\n"
        f"### {pin} Your message details:\n\n"
        f"{message}\n\n"
        "In the meantime, you can continue using our tools to:\n\n"
        "* Compress files\n"
        "* Convert file formats/extensions\n"
        "* Split or merge PDFs\n"
        "* And more useful features\n\n"
        "We usually respond within **24 hours**.\n\n"
        "---\n\n"
        "If your query is urgent, feel free to reply directly to this email.\n\n"
        f"Thanks for contacting us {smile}\n\n"
        "**With regards,**\n"
        "Sohail\n\n"
        "---\n"
    )


def _contact_template_html(name: str, message: str) -> str:
    n, m = html.escape(name), html.escape(message).replace("\n", "<br />\n")
    wave, pin, smile = "\U0001f64c", "\U0001f4cc", "\U0001f60a"
    return f"""<!DOCTYPE html>
<html><body style="font-family:system-ui,sans-serif;line-height:1.55;color:#0f172a;">
<p>Hi {n},</p>
<p>Thanks for reaching out to us! {wave}</p>
<p>We've received your message and our team will get back to you as soon as possible.</p>
<h3>{pin} Your message details:</h3>
<p style="white-space:pre-wrap;border-left:3px solid #6366f1;padding-left:12px;">{m}</p>
<p>In the meantime, you can continue using our tools to:</p>
<ul>
<li>Compress files</li>
<li>Convert file formats/extensions</li>
<li>Split or merge PDFs</li>
<li>And more useful features</li>
</ul>
<p>We usually respond within <strong>24 hours</strong>.</p>
<hr />
<p>If your query is urgent, feel free to reply directly to this email.</p>
<p>Thanks for contacting us {smile}</p>
<p><strong>With regards,</strong><br />Sohail</p>
<hr />
</body></html>"""


@app.get("/api/health")
def health():
    return JSONResponse({"ok": True})


@app.post("/api/contact")
def contact_submit(payload: ContactPayload):
    try:
        from_header = _contact_from_header()
        inbox = _contact_inbox()
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    subj = payload.subject or "File Mitra"
    team = EmailMessage()
    team["Subject"] = f"[Contact] {subj}"
    team["From"] = from_header
    team["To"] = inbox
    team["Reply-To"] = str(payload.email)
    team.set_content(
        f"Name: {payload.name}\nEmail: {payload.email}\nSubject: {subj}\n\n{payload.message}\n",
        charset="utf-8",
    )

    user = EmailMessage()
    user["Subject"] = "Thanks for contacting File Mitra"
    user["From"] = from_header
    user["To"] = str(payload.email)
    user["Reply-To"] = inbox
    user.set_content(_contact_template_plain(payload.name, payload.message), charset="utf-8")
    user.add_alternative(_contact_template_html(payload.name, payload.message), subtype="html", charset="utf-8")

    try:
        _send_smtp([team, user])
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except OSError as exc:
        raise HTTPException(status_code=502, detail=f"Mail delivery failed: {exc!s}") from exc
    except smtplib.SMTPException as exc:
        raise HTTPException(status_code=502, detail=f"Mail delivery failed: {exc!s}") from exc

    return JSONResponse({"ok": True})


@app.post("/api/convert/image")
def convert_image(
    file: UploadFile = File(...),
    format: Literal["jpeg", "png", "webp", "bmp"] = Form(...),
    quality: int = Form(92),
):
    img = _read_image(file)
    out = io.BytesIO()
    target = format.upper()
    save_args: dict[str, object] = {}

    if target == "JPEG":
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        save_args["quality"] = max(20, min(100, int(quality)))
        save_args["optimize"] = True
    elif target == "WEBP":
        save_args["quality"] = max(20, min(100, int(quality)))
        save_args["method"] = 6

    img.save(out, format=target, **save_args)
    base = os.path.splitext(file.filename or "image")[0]
    filename = f"{_safe_name(base, 'image')}.{format}"
    return _stream_bytes(out.getvalue(), filename, "application/octet-stream")


@app.post("/api/compress/image")
def compress_image(
    file: UploadFile = File(...),
    quality: int = Form(80),
    max_width: int = Form(1920),
):
    img = _read_image(file)
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")

    width, height = img.size
    if width > max_width > 0:
        new_height = int((height * max_width) / width)
        img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)

    out = io.BytesIO()
    img.save(out, format="JPEG", quality=max(20, min(95, quality)), optimize=True)
    base = os.path.splitext(file.filename or "image")[0]
    filename = f"{_safe_name(base, 'image')}-compressed.jpg"
    return _stream_bytes(out.getvalue(), filename, "application/octet-stream")


@app.post("/api/pdf/make")
def make_pdf(
    files: list[UploadFile] = File(...),
    page_size: str = Form("a4"),
    orientation: str = Form("portrait"),
    margin: int = Form(24),
):
    if not files:
        raise HTTPException(status_code=400, detail="No images provided.")

    images: list[Image.Image] = []
    for f in files:
        img = _read_image(f)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        images.append(img)

    first, rest = images[0], images[1:]
    out = io.BytesIO()
    first.save(out, format="PDF", save_all=True, append_images=rest)
    filename = "document.pdf"
    # Keep args consumed for frontend compatibility even if not layout-applied now
    _ = page_size, orientation, margin
    return _stream_bytes(out.getvalue(), filename, "application/pdf")


@app.post("/api/pdf/merge")
def merge_pdf(files: list[UploadFile] = File(...)):
    if len(files) < 2:
        raise HTTPException(
            status_code=400, detail="At least 2 PDF files are required."
        )

    writer = PdfWriter()
    for file in files:
        raw = file.file.read()
        if not raw:
            continue
        reader = PdfReader(io.BytesIO(raw))
        for page in reader.pages:
            writer.add_page(page)

    if len(writer.pages) == 0:
        raise HTTPException(status_code=400, detail="No valid PDF pages found.")

    out = io.BytesIO()
    writer.write(out)
    return _stream_bytes(out.getvalue(), "merged.pdf", "application/pdf")


@app.post("/api/compress/pdf")
def compress_pdf(
    file: UploadFile = File(...),
    level: Literal["low", "medium", "high"] = Form("medium"),
):
    raw = file.file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Empty PDF file.")

    try:
        doc = fitz.open(stream=raw, filetype="pdf")
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=400, detail="Invalid PDF file.") from exc

    # PyMuPDF built-in save optimization settings.
    deflate_images = level in ("medium", "high")
    deflate_fonts = level == "high"
    out = io.BytesIO()
    doc.save(
        out,
        garbage=4,
        deflate=True,
        deflate_images=deflate_images,
        deflate_fonts=deflate_fonts,
        clean=True,
    )
    doc.close()
    return _stream_bytes(out.getvalue(), "compressed.pdf", "application/pdf")


@app.post("/api/pdf/convert")
def convert_pdf(
    file: UploadFile = File(...),
    output: Literal["xlsx", "docx"] = Form(...),
    mode: Literal["tables", "text"] = Form("text"),
):
    raw = file.file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Empty PDF file.")
    try:
        doc = fitz.open(stream=raw, filetype="pdf")
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=400, detail="Invalid PDF file.") from exc

    pages_text: list[str] = []
    for page in doc:
        pages_text.append(page.get_text("text") or "")
    doc.close()

    if output == "docx":
        word = Document()
        for i, content in enumerate(pages_text, start=1):
            word.add_heading(f"Page {i}", level=2)
            word.add_paragraph(content.strip() or "(No extractable text)")
        out = io.BytesIO()
        word.save(out)
        return _stream_bytes(
            out.getvalue(),
            "converted.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Extracted"
    sheet.append(["Page", "Line", "Text"])
    for page_no, content in enumerate(pages_text, start=1):
        lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
        if not lines:
            sheet.append([page_no, 1, "(No extractable text)"])
            continue
        for i, line in enumerate(lines, start=1):
            sheet.append([page_no, i, line])

    # Auto-size column widths based on extracted content to keep XLSX readable.
    for col in range(1, sheet.max_column + 1):
        max_len = 0
        for row in range(1, sheet.max_row + 1):
            value = sheet.cell(row=row, column=col).value
            if value is None:
                continue
            text = str(value)
            cell_len = max(len(part) for part in text.splitlines()) if text else 0
            if cell_len > max_len:
                max_len = cell_len
        # Keep practical bounds so a single long line does not break layout.
        sheet.column_dimensions[get_column_letter(col)].width = min(max(max_len + 2, 8), 70)

    # Make extracted text easier to scan when long lines wrap.
    for row in range(2, sheet.max_row + 1):
        sheet.cell(row=row, column=3).alignment = Alignment(vertical="top", wrap_text=True)

    _ = mode
    out = io.BytesIO()
    workbook.save(out)
    return _stream_bytes(
        out.getvalue(),
        "converted.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.post("/api/ocr/extract")
def ocr_extract(
    file: UploadFile = File(...),
    language: str = Form("eng"),
    output: Literal["txt", "pdf", "docx"] = Form("txt"),
):
    name = (file.filename or "").lower()
    raw = file.file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Empty file.")

    # Word (.docx): extract embedded text (no OCR — fast)
    if name.endswith(".docx"):
        try:
            extracted = _extract_text_from_docx(raw)
        except Exception as exc:  # pragma: no cover
            raise HTTPException(
                status_code=400, detail="Invalid Word (.docx) file."
            ) from exc
        if not extracted:
            extracted = "No text found in document."
        if output == "txt":
            return _stream_bytes(
                extracted.encode("utf-8"), "ocr-output.txt", "text/plain"
            )
        if output == "docx":
            doc = Document()
            doc.add_heading("Extracted Text", level=1)
            doc.add_paragraph(extracted)
            out = io.BytesIO()
            doc.save(out)
            return _stream_bytes(
                out.getvalue(),
                "extracted.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        rect = fitz.Rect(40, 40, 555, 802)
        page.insert_textbox(rect, extracted, fontsize=11)
        out = io.BytesIO()
        pdf_doc.save(out)
        pdf_doc.close()
        return _stream_bytes(out.getvalue(), "extracted.pdf", "application/pdf")

    text_result = ""
    images_for_ocr: list[Image.Image] = []
    pdf_truncation_note = ""

    ocr_dpi = 150
    max_ocr_pages = 15
    max_embedded_pages = 50

    if name.endswith(".pdf") or _is_pdf_bytes(raw):
        try:
            pdf = fitz.open(stream=raw, filetype="pdf")
        except Exception as exc:  # pragma: no cover
            raise HTTPException(status_code=400, detail="Invalid PDF file.") from exc
        try:
            page_count = pdf.page_count
            embed_n = min(page_count, max_embedded_pages)
            chunks: list[str] = []
            for i in range(embed_n):
                page = pdf.load_page(i)
                t = (page.get_text("text") or "").strip()
                if t:
                    chunks.append(t)
            embedded = "\n\n".join(chunks).strip()

            if len(embedded) >= _MIN_EMBEDDED_PDF_TEXT:
                text_result = embedded
                if page_count > embed_n:
                    pdf_truncation_note = f"\n\n[Note: Text from first {embed_n} of {page_count} pages only.]\n"
            else:
                for i in range(min(page_count, max_ocr_pages)):
                    page = pdf.load_page(i)
                    mat = fitz.Matrix(ocr_dpi / 72, ocr_dpi / 72)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    images_for_ocr.append(_resize_for_ocr(img))
                if page_count > max_ocr_pages:
                    pdf_truncation_note = f"\n\n[Note: Only first {max_ocr_pages} of {page_count} pages were OCR'd for speed.]\n"
        finally:
            pdf.close()
    else:
        try:
            img = Image.open(io.BytesIO(raw))
            img.load()
            images_for_ocr.append(_resize_for_ocr(_image_to_rgb(img)))
        except Exception as exc:  # pragma: no cover
            raise HTTPException(
                status_code=400,
                detail="Unsupported file. Use image (JPG/PNG/WEBP), PDF, or .docx.",
            ) from exc

    for idx, image in enumerate(images_for_ocr):
        try:
            text_result += _ocr_image_pil(image, language) + "\n"
        except HTTPException:
            raise
        except Exception as exc:  # pragma: no cover
            raise HTTPException(
                status_code=500, detail=f"OCR failed on page/slice {idx + 1}: {exc!s}"
            ) from exc

    cleaned = (
        text_result.strip() + pdf_truncation_note
    ).strip() or "No OCR text extracted."

    if output == "txt":
        return _stream_bytes(cleaned.encode("utf-8"), "ocr-output.txt", "text/plain")

    if output == "docx":
        doc = Document()
        doc.add_heading("OCR Output", level=1)
        doc.add_paragraph(cleaned)
        out = io.BytesIO()
        doc.save(out)
        return _stream_bytes(
            out.getvalue(),
            "ocr-output.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    rect = fitz.Rect(40, 40, 555, 802)
    page.insert_textbox(rect, cleaned, fontsize=11)
    out = io.BytesIO()
    pdf_doc.save(out)
    pdf_doc.close()
    return _stream_bytes(out.getvalue(), "ocr-output.pdf", "application/pdf")


@app.post("/generate-image", response_model=ImageGenerationResponse)
@app.post("/api/generate-image", response_model=ImageGenerationResponse)
async def generate_image(payload: ImageGenerationRequest, request: Request):
    client_ip = (
        (request.headers.get("x-forwarded-for", "").split(",")[0].strip())
        or (request.client.host if request.client else "")
        or "unknown"
    )
    rate_limiter.check(client_ip)

    if payload.prompt.count("\n") > 25:
        raise HTTPException(status_code=400, detail="Prompt format looks abusive. Keep it concise.")

    cleaned_prompt = enhance_prompt(payload.prompt)
    if not cleaned_prompt:
        raise HTTPException(status_code=400, detail="Prompt is required.")

    _reload_ai_env()
    try:
        provider = build_image_provider()
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    logger.info("Image generation request received from ip=%s", client_ip)
    image_bytes = await provider.generate_image(cleaned_prompt)
    filename = f"{uuid.uuid4().hex}.png"
    output_path = generated_images_dir / filename
    output_path.write_bytes(image_bytes)
    image_url = f"{GENERATED_IMAGES_DIR}/{filename}"

    return ImageGenerationResponse(
        message="Ye raha tera generated image bhai 🔥",
        image_url=image_url,
    )


# backend/main.py → parent dir is frontend app root (where package.json / dist/ live)
# _BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
# DIST_PATH = os.environ.get("FRONTEND_DIST_PATH", os.path.join(_BASE, "dist"))
# _ASSETS_DIR = os.path.join(DIST_PATH, "assets")
# _INDEX_HTML = os.path.join(DIST_PATH, "index.html")

# if os.path.isdir(_ASSETS_DIR):
#     app.mount("/assets", StaticFiles(directory=_ASSETS_DIR), name="assets")


# @app.get("/")
# def serve_react():
#     if not os.path.isfile(_INDEX_HTML):
#         return JSONResponse(
#             status_code=503,
#             content={
#                 "detail": "Frontend build not found. Run `npm run build` in the frontend root.",
#                 "expected_index": _INDEX_HTML,
#             },
#         )
#     return FileResponse(_INDEX_HTML)

@app.get("/")
def root():
    return {"status": "backend running"}
